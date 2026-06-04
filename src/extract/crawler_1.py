"""
Montana State Government Website Crawler
Crawls mt.gov websites and converts content to AI-friendly markdown
Supports HTML pages, PDF documents, and DOCX files
"""

import os
import re
import requests
from bs4 import BeautifulSoup
from markdownify import markdownify as md
from urllib.parse import urljoin, urlparse
import time
from pathlib import Path
from PyPDF2 import PdfReader
import io
from docx import Document
import json


class StateAgencyCrawler:
    def __init__(self, base_url, output_dir="knowledge", agency_name="Montana State Agency", subfolder=None):
        self.base_url = base_url
        self.agency_name = agency_name
        
        # If subfolder is specified, create agency-specific directory
        if subfolder:
            self.output_dir = Path(output_dir) / subfolder
        else:
            self.output_dir = Path(output_dir)
            
        self.visited_urls = set()
        self.to_visit = set([base_url])
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'
        })
        self.pdf_count = 0
        self.docx_count = 0
        self.domain = urlparse(base_url).netloc
        
        # Inclusion list - only crawl these file types
        self.allowed_extensions = {'.html', '.htm', '.pdf', '.docx'}
        
        # Navigation network tracking
        self.navigation_network = {
            'nodes': {},      # url -> {title, type, file_path, metadata}
            'edges': []       # list of {source, target, link_text, context}
        }
        
    def is_valid_url(self, url):
        """Check if URL belongs to the agency's domain and matches inclusion list"""
        parsed = urlparse(url)
        
        # Check if URL belongs to the same domain as base_url
        if not (parsed.netloc == self.domain or parsed.netloc == ''):
            return False
        
        # Get file extension
        path_lower = parsed.path.lower()
        
        # If it has an extension, check if it's in our allowed list
        if any(path_lower.endswith(ext) for ext in ['.pdf', '.docx', '.doc']):
            return True
        
        # Allow HTML pages (with or without .html/.htm extension)
        # URLs without extensions are assumed to be HTML pages
        if any(path_lower.endswith(ext) for ext in ['.html', '.htm']):
            return True
        
        # If no extension or ends with '/', assume it's an HTML page
        if '.' not in Path(parsed.path).name or path_lower.endswith('/'):
            return True
        
        return False
    
    def normalize_url(self, url):
        """Normalize URL for comparison"""
        parsed = urlparse(url)
        # Remove fragments and trailing slashes
        normalized = f"{parsed.scheme}://{parsed.netloc}{parsed.path.rstrip('/')}"
        if parsed.query:
            normalized += f"?{parsed.query}"
        return normalized
    
    def get_output_path(self, url):
        """Convert URL to file path"""
        parsed = urlparse(url)
        path = parsed.path.strip('/')
        
        if not path:
            path = 'index'
        
        # Clean up path for filename
        path = path.replace('/', '_')
        path = re.sub(r'[^\w\-_]', '_', path)
        
        return self.output_dir / f"{path}.md"
    
    def extract_links(self, soup, current_url):
        """Extract all valid links from page, including PDFs and DOCX files"""
        links = set()
        pdf_links = set()
        docx_links = set()
        
        for link in soup.find_all('a', href=True):
            href = link['href']
            absolute_url = urljoin(current_url, href)
            link_text = link.get_text().strip() or href
            
            # Get context (surrounding text)
            context = ''
            if link.parent:
                context = link.parent.get_text()[:200].strip()
            
            # Check if it's a PDF
            if absolute_url.lower().endswith('.pdf'):
                if self.is_valid_url(absolute_url):
                    pdf_links.add(absolute_url)
                    # Track edge to PDF
                    self.navigation_network['edges'].append({
                        'source': current_url,
                        'target': absolute_url,
                        'link_text': link_text,
                        'context': context,
                        'type': 'pdf_link'
                    })
            # Check if it's a DOCX
            elif absolute_url.lower().endswith(('.docx', '.doc')):
                if self.is_valid_url(absolute_url):
                    docx_links.add(absolute_url)
                    # Track edge to DOCX
                    self.navigation_network['edges'].append({
                        'source': current_url,
                        'target': absolute_url,
                        'link_text': link_text,
                        'context': context,
                        'type': 'docx_link'
                    })
            # Only include same-domain links (HTML pages)
            elif self.is_valid_url(absolute_url):
                normalized = self.normalize_url(absolute_url)
                if normalized.startswith(self.base_url.rstrip('/')):
                    links.add(normalized)
                    # Track edge to HTML page
                    self.navigation_network['edges'].append({
                        'source': current_url,
                        'target': normalized,
                        'link_text': link_text,
                        'context': context,
                        'type': 'hyperlink'
                    })
        
        return links, pdf_links, docx_links
    
    def clean_content(self, soup):
        """Remove navigation, scripts, and other non-content elements"""
        # Remove script and style elements
        for element in soup(['script', 'style', 'nav', 'footer', 'header']):
            element.decompose()
        
        # Remove common navigation/sidebar classes
        for element in soup.find_all(class_=re.compile(r'nav|sidebar|footer|header|menu', re.I)):
            element.decompose()
        
        # Remove common navigation IDs
        for element in soup.find_all(id=re.compile(r'nav|sidebar|footer|header|menu', re.I)):
            element.decompose()
            
        return soup
    
    def html_to_markdown(self, html, url):
        """Convert HTML to clean markdown"""
        soup = BeautifulSoup(html, 'html.parser')
        
        # Clean the content
        soup = self.clean_content(soup)
        
        # Get the main content area if it exists
        main_content = soup.find('main') or soup.find('article') or soup.find(id='content') or soup.find(class_=re.compile(r'content|main', re.I)) or soup
        
        # Get page title
        title = soup.find('title')
        title_text = title.get_text().strip() if title else "Untitled"
        
        # Convert to markdown
        markdown_content = md(str(main_content), heading_style="ATX", bullets="-")
        
        # Create header with metadata
        header = f"""---
source: {url}
title: {title_text}
crawled: {time.strftime('%Y-%m-%d %H:%M:%S')}
---

# {title_text}

"""
        
        # Clean up markdown (remove excessive newlines)
        markdown_content = re.sub(r'\n{3,}', '\n\n', markdown_content)
        
        return header + markdown_content
    
    def extract_pdf_content(self, url):
        """Download and extract text from PDF"""
        try:
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            
            # Read PDF from bytes
            pdf_file = io.BytesIO(response.content)
            pdf_reader = PdfReader(pdf_file)
            
            # Extract text from all pages
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"## Page {page_num}\n\n{text}")
            
            # Get PDF metadata
            metadata = pdf_reader.metadata
            title = metadata.title if metadata and metadata.title else Path(url).stem
            
            # Create markdown header
            header = f"""---
source: {url}
title: {title}
type: PDF
pages: {len(pdf_reader.pages)}
crawled: {time.strftime('%Y-%m-%d %H:%M:%S')}
---

# {title}

"""
            
            return header + "\n\n".join(text_content)
            
        except Exception as e:
            print(f"  ✗ Error extracting PDF {url}: {str(e)}")
            return None
    
    def process_pdf(self, url):
        """Process a PDF file and save as markdown"""
        print(f"Extracting PDF: {url}")
        
        markdown_content = self.extract_pdf_content(url)
        if not markdown_content:
            return
        
        # Create filename from URL
        parsed = urlparse(url)
        filename = Path(parsed.path).stem
        filename = re.sub(r'[^\w\-_]', '_', filename)
        output_path = self.output_dir / f"pdf_{filename}.md"
        
        # Save to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        # Add to navigation network
        parsed = urlparse(url)
        filename = Path(parsed.path).stem
        self.navigation_network['nodes'][url] = {
            'title': filename,
            'type': 'pdf_document',
            'file_path': str(output_path.relative_to(self.output_dir)),
            'url': url,
            'crawled': time.strftime('%Y-%m-%d %H:%M:%S')
        }
        
        self.pdf_count += 1
        print(f"  → Saved PDF to: {output_path}")
    
    def extract_docx_content(self, url):
        """Download and extract text from DOCX file"""
        try:
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            
            # Read DOCX from bytes
            docx_file = io.BytesIO(response.content)
            doc = Document(docx_file)
            
            # Extract content with structure
            markdown_parts = []
            
            # Process each paragraph
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Check if it's a heading based on style
                if para.style.name.startswith('Heading'):
                    # Extract heading level
                    try:
                        level = int(para.style.name.replace('Heading ', ''))
                        markdown_parts.append(f"{'#' * level} {text}")
                    except:
                        markdown_parts.append(f"## {text}")
                else:
                    markdown_parts.append(text)
                    
                markdown_parts.append("")  # Add blank line
            
            # Process tables
            for table_idx, table in enumerate(doc.tables, 1):
                markdown_parts.append(f"\n### Table {table_idx}\n")
                
                for row_idx, row in enumerate(table.rows):
                    # Extract cell text
                    cells = [cell.text.strip() for cell in row.cells]
                    markdown_parts.append("| " + " | ".join(cells) + " |")
                    
                    # Add header separator after first row
                    if row_idx == 0:
                        markdown_parts.append("|" + "|".join([" --- " for _ in cells]) + "|")
                
                markdown_parts.append("")  # Add blank line after table
            
            # Get document title (from filename)
            title = Path(url).stem
            
            # Create markdown header
            header = f"""---
source: {url}
title: {title}
type: DOCX
crawled: {time.strftime('%Y-%m-%d %H:%M:%S')}
---

# {title}

"""
            
            content = "\n".join(markdown_parts)
            # Clean up excessive newlines
            content = re.sub(r'\n{3,}', '\n\n', content)
            
            return header + content
            
        except Exception as e:
            print(f"  ✗ Error extracting DOCX {url}: {str(e)}")
            return None
    
    def process_docx(self, url):
        """Process a DOCX file and save as markdown"""
        print(f"Extracting DOCX: {url}")
        
        markdown_content = self.extract_docx_content(url)
        if not markdown_content:
            return
        
        # Create filename from URL
        parsed = urlparse(url)
        filename = Path(parsed.path).stem
        filename = re.sub(r'[^\w\-_]', '_', filename)
        output_path = self.output_dir / f"_docs_{filename}.md"
        
        # Save to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        # Add to navigation network
        parsed = urlparse(url)
        filename = Path(parsed.path).stem
        self.navigation_network['nodes'][url] = {
            'title': filename,
            'type': 'docx_document',
            'file_path': str(output_path.relative_to(self.output_dir)),
            'url': url,
            'crawled': time.strftime('%Y-%m-%d %H:%M:%S')
        }
        
        self.docx_count += 1
        print(f"  → Saved DOCX to: {output_path}")
    
    def crawl_page(self, url):
        """Crawl a single page"""
        print(f"Crawling: {url}")
        
        try:
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            
            # Convert to markdown
            markdown_content = self.html_to_markdown(response.text, url)
            
            # Save to file
            output_path = self.get_output_path(url)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            print(f"  → Saved to: {output_path}")
            
            # Extract links for further crawling
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Add node to navigation network
            title = soup.find('title')
            title_text = title.get_text().strip() if title else "Untitled"
            self.navigation_network['nodes'][url] = {
                'title': title_text,
                'type': 'html_page',
                'file_path': str(output_path.relative_to(self.output_dir)),
                'url': url,
                'crawled': time.strftime('%Y-%m-%d %H:%M:%S')
            }
            new_links, pdf_links, docx_links = self.extract_links(soup, url)
            
            # Process any PDFs found on this page
            for pdf_url in pdf_links:
                if pdf_url not in self.visited_urls:
                    self.visited_urls.add(pdf_url)
                    self.process_pdf(pdf_url)
            
            # Process any DOCX files found on this page
            for docx_url in docx_links:
                if docx_url not in self.visited_urls:
                    self.visited_urls.add(docx_url)
                    self.process_docx(docx_url)
            
            return new_links
            
        except Exception as e:
            print(f"  ✗ Error crawling {url}: {str(e)}")
            return set()
    
    def crawl(self):
        """Main crawl loop"""
        print(f"\n{'='*80}")
        print(f"Starting crawl of {self.agency_name}")
        print(f"URL: {self.base_url}")
        print(f"Output directory: {self.output_dir.absolute()}")
        print(f"{'='*80}\n")
        
        # Create output directory
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        while self.to_visit:
            url = self.to_visit.pop()
            
            if url in self.visited_urls:
                continue
            
            self.visited_urls.add(url)
            
            # Crawl the page
            new_links = self.crawl_page(url)
            
            # Add new links to queue
            unvisited_links = new_links - self.visited_urls
            self.to_visit.update(unvisited_links)
            
            # Be polite - small delay between requests
            time.sleep(0.5)
        
        # Save navigation network
        network_file = self.output_dir / 'navigation_network.json'
        network_data = {
            'agency': self.agency_name,
            'base_url': self.base_url,
            'created': time.strftime('%Y-%m-%d %H:%M:%S'),
            'statistics': {
                'total_nodes': len(self.navigation_network['nodes']),
                'total_edges': len(self.navigation_network['edges']),
                'html_pages': sum(1 for n in self.navigation_network['nodes'].values() if n['type'] == 'html_page'),
                'pdf_documents': sum(1 for n in self.navigation_network['nodes'].values() if n['type'] == 'pdf_document'),
                'docx_documents': sum(1 for n in self.navigation_network['nodes'].values() if n['type'] == 'docx_document')
            },
            'nodes': self.navigation_network['nodes'],
            'edges': self.navigation_network['edges']
        }
        
        with open(network_file, 'w', encoding='utf-8') as f:
            json.dump(network_data, f, indent=2, ensure_ascii=False)
        
        print(f"\n{'='*80}")
        print(f"✓ Crawl complete for {self.agency_name}!")
        print(f"  Pages crawled: {len(self.visited_urls)}")
        print(f"  PDFs extracted: {self.pdf_count}")
        print(f"  DOCX files extracted: {self.docx_count}")
        print(f"  Files created: {len(list(self.output_dir.glob('*.md')))}")
        print(f"  Navigation network saved: {network_file}")
        print(f"  Network nodes: {len(self.navigation_network['nodes'])}")
        print(f"  Network edges: {len(self.navigation_network['edges'])}")
        print(f"{'='*80}\n")
        
        return {
            'agency': self.agency_name,
            'pages': len(self.visited_urls),
            'pdfs': self.pdf_count,
            'docx': self.docx_count,
            'files': len(list(self.output_dir.glob('*.md'))),
            'network_nodes': len(self.navigation_network['nodes']),
            'network_edges': len(self.navigation_network['edges'])
        }


# This file is a library module.
# Use 2_refresh.py to crawl agencies from agencies.md
# Or import StateAgencyCrawler to use programmatically:
#
#   from src.extract.crawler import StateAgencyCrawler
#   crawler = StateAgencyCrawler(base_url, agency_name, subfolder)
#   crawler.crawl()
