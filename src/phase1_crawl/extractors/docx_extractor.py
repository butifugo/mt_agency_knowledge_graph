"""
DOCX Extractor
Extracts text content from DOCX documents
"""

import io
import re
from pathlib import Path
from docx import Document as DocxDocument
from typing import Optional


class DocxExtractor:
    """Extract text from DOCX documents"""
    
    @staticmethod
    def extract(docx_bytes: bytes, url: str, timestamp: str) -> Optional[str]:
        """
        Extract text from DOCX and convert to markdown
        
        Args:
            docx_bytes: DOCX file bytes
            url: Source URL
            timestamp: Crawl timestamp
            
        Returns:
            Markdown content or None if extraction fails
        """
        try:
            # Read DOCX from bytes
            docx_file = io.BytesIO(docx_bytes)
            doc = DocxDocument(docx_file)
            
            # Extract content with structure
            markdown_parts = []
            
            # Process each paragraph
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Check if it's a heading based on style
                style_name = para.style.name if para.style else ""
                if style_name and style_name.startswith('Heading'):
                    # Extract heading level
                    try:
                        level = int(style_name.replace('Heading ', ''))
                        markdown_parts.append(f"{'#' * level} {text}")
                    except:
                        markdown_parts.append(f"## {text}")
                else:
                    markdown_parts.append(text)
                    
                markdown_parts.append("")  # Add blank line
            
            # Process tables
            for table_idx, table in enumerate(doc.tables, 1):
                markdown_parts.append(f"\n### Table {table_idx}\n")
                
                for row_idx, row in enumerate(table.rows):
                    # Extract cell text
                    cells = [cell.text.strip() for cell in row.cells]
                    markdown_parts.append("| " + " | ".join(cells) + " |")
                    
                    # Add header separator after first row
                    if row_idx == 0:
                        markdown_parts.append("|" + "|".join([" --- " for _ in cells]) + "|")
                
                markdown_parts.append("")  # Add blank line after table
            
            # Get document title (from filename)
            title = Path(url).stem
            
            # Create markdown header
            header = f"""---
source: {url}
title: {title}
type: docx
crawled: {timestamp}
---

# {title}

"""
            
            content = "\n".join(markdown_parts)
            # Clean up excessive newlines
            content = re.sub(r'\n{3,}', '\n\n', content)
            
            return header + content
            
        except Exception as e:
            print(f"  ✗ Error extracting DOCX: {str(e)}")
            return None
